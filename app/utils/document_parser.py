"""
Document parser for extracting text from various file formats.
Supports PDF, DOCX, and TXT files with encoding detection.
"""

from io import BytesIO
from pathlib import Path
from typing import BinaryIO, Union

import chardet
from docx import Document as DocxDocument
from PyPDF2 import PdfReader

from app.core.exceptions import DocumentProcessingError, UnsupportedFileTypeError
from app.core.logging import get_logger
from app.models.enums import FileType

logger = get_logger(__name__)


class DocumentParser:
    """
    Parser for extracting text content from documents.
    
    Supports:
    - PDF files (using PyPDF2)
    - DOCX files (using python-docx)
    - TXT files (with encoding detection)
    """

    def __init__(self, allowed_types: list[str] | None = None):
        """
        Initialize the document parser.
        
        Args:
            allowed_types: List of allowed file extensions (without dot).
                          Defaults to ['pdf', 'docx', 'txt'].
        """
        self.allowed_types = allowed_types or ["pdf", "docx", "txt"]
        logger.info("document_parser_initialized", allowed_types=self.allowed_types)

    def parse(
        self,
        file_content: Union[bytes, BinaryIO],
        filename: str,
    ) -> str:
        """
        Parse document and extract text content.
        
        Args:
            file_content: File content as bytes or file-like object.
            filename: Original filename to determine file type.
            
        Returns:
            Extracted text content.
            
        Raises:
            UnsupportedFileTypeError: If file type is not supported.
            DocumentProcessingError: If parsing fails.
        """
        # Get file extension
        extension = Path(filename).suffix.lower().lstrip(".")
        
        if extension not in self.allowed_types:
            raise UnsupportedFileTypeError(
                filename=filename,
                file_type=extension,
                allowed_types=self.allowed_types,
            )

        # Convert to bytes if needed
        if hasattr(file_content, "read"):
            content_bytes = file_content.read()
        else:
            content_bytes = file_content

        logger.debug(
            "parsing_document",
            filename=filename,
            extension=extension,
            size_bytes=len(content_bytes),
        )

        try:
            file_type = FileType.from_extension(extension)
            
            if file_type == FileType.PDF:
                return self._parse_pdf(content_bytes, filename)
            elif file_type == FileType.DOCX:
                return self._parse_docx(content_bytes, filename)
            elif file_type == FileType.TXT:
                return self._parse_txt(content_bytes, filename)
            else:
                raise UnsupportedFileTypeError(
                    filename=filename,
                    file_type=extension,
                    allowed_types=self.allowed_types,
                )
        except (UnsupportedFileTypeError, DocumentProcessingError):
            raise
        except Exception as e:
            logger.error("document_parsing_failed", filename=filename, error=str(e))
            raise DocumentProcessingError(filename=filename, reason=str(e))

    def _parse_pdf(self, content: bytes, filename: str) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PdfReader(BytesIO(content))
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    
            text = "\n\n".join(text_parts)
            
            if not text.strip():
                logger.warning("pdf_no_text_extracted", filename=filename)
                
            logger.info(
                "pdf_parsed",
                filename=filename,
                pages=len(pdf_reader.pages),
                chars=len(text),
            )
            return text
            
        except Exception as e:
            raise DocumentProcessingError(filename=filename, reason=f"PDF parsing error: {e}")

    def _parse_docx(self, content: bytes, filename: str) -> str:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
                    
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
                        
            text = "\n\n".join(text_parts)
            
            logger.info(
                "docx_parsed",
                filename=filename,
                paragraphs=len(doc.paragraphs),
                tables=len(doc.tables),
                chars=len(text),
            )
            return text
            
        except Exception as e:
            raise DocumentProcessingError(filename=filename, reason=f"DOCX parsing error: {e}")

    def _parse_txt(self, content: bytes, filename: str) -> str:
        """Extract text from TXT file with encoding detection."""
        try:
            # Detect encoding
            detected = chardet.detect(content)
            encoding = detected.get("encoding", "utf-8") or "utf-8"
            confidence = detected.get("confidence", 0)
            
            logger.debug(
                "encoding_detected",
                filename=filename,
                encoding=encoding,
                confidence=confidence,
            )
            
            # Try detected encoding first, fall back to utf-8
            try:
                text = content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                text = content.decode("utf-8", errors="replace")
                logger.warning("encoding_fallback", filename=filename, original_encoding=encoding)
                
            logger.info("txt_parsed", filename=filename, chars=len(text))
            return text
            
        except Exception as e:
            raise DocumentProcessingError(filename=filename, reason=f"TXT parsing error: {e}")

    def get_file_type(self, filename: str) -> FileType:
        """
        Get FileType from filename.
        
        Args:
            filename: Filename to check.
            
        Returns:
            FileType enum value.
            
        Raises:
            UnsupportedFileTypeError: If extension is not supported.
        """
        extension = Path(filename).suffix.lower().lstrip(".")
        
        if extension not in self.allowed_types:
            raise UnsupportedFileTypeError(
                filename=filename,
                file_type=extension,
                allowed_types=self.allowed_types,
            )
            
        return FileType.from_extension(extension)
